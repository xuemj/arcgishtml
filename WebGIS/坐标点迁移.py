import os
from tkinter import filedialog
from copy import deepcopy
from docx import Document
import tkinter as tk
from tkinter import messagebox
import time
from docx.enum.table import WD_TABLE_ALIGNMENT 

root = tk.Tk()
root.withdraw()

messagebox.showinfo('提示','请选择数据文件！')

folderpath = filedialog.askopenfilename() 

arr1 = []
arr2 = []

doc = Document(folderpath)

for table in doc.tables :
    if table.cell(0,0).text == '界  址  标  示' :
        for i in range(len(table.rows)) :
            if 'J' in table.cell(i,0).text :
                ss = table.cell(i,0).text.strip('J')
                if int(ss) > 56 :
                    arr1.append(table.cell(i,0).text)
                    arr2.append(table.cell(i,7).text)
print(arr1)
print(arr2)
messagebox.showinfo('提示','请选择迁移后的文件！')

folderpath1 = filedialog.askopenfilename()


count =  len(arr1)//27 + 1

doc1 = Document(folderpath1)
for i in range(0,count) :
    table = doc1.tables[2]  
    new_table = deepcopy(table._tbl)
    p = doc1.paragraphs[21]
    paragraph = p.insert_paragraph_before()
    paragraph._p.addnext(new_table)
    doc1.save(folderpath1)

table1 = doc1.tables[5]
table1.cell(30,6).text = arr2[0]

num_bj = len(arr1)
num = num_bj - 1
ar_num = 0
for i in range(6,6+count) :
    table2 = doc1.tables[i]
    for j in range(3,30) :
        ar_num += 1
        if num > 27 :
            table2.cell(j,0).text = arr1[ar_num]
            table2.cell(j,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table2.cell(j,6).text = arr2[ar_num]
            num = num -27
        else :
            print(j,ar_num)
            if ar_num < num_bj :
                table2.cell(j,0).text = arr1[ar_num]
                table2.cell(j,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table2.cell(j,6).text = arr2[ar_num]
            if ar_num == num_bj :
                table2.cell(j,0).text = 'J1'
                table2.cell(j,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            if ar_num > num_bj+1 :
                table2.cell(j,0).text = ''
                table2.cell(j,4).text = ''
                table2.cell(j,6).text = ''
                table2.cell(j,7).text = ''
                table2.cell(j,17).text = ''
doc1.save(folderpath1)
messagebox.showinfo('提示','填充完成！')