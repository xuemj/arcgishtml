import os
from tkinter import filedialog
from docx import Document
import tkinter as tk
from tkinter import messagebox

root = tk.Tk()
root.withdraw()

messagebox.showinfo('提示','请选择文件夹！')

Folderpath = filedialog.askdirectory() #获得选择好的文件夹


# 获取filepath文件夹下的所有的文件
def getfilelist(filepath):
    filelist =  os.listdir(filepath)
    files = []
    for i in range(len(filelist)):
        child = os.path.join('%s/%s'%(filepath, filelist[i]))
        if os.path.isdir(child):
            files.extend(getfilelist(child))
        else:
            files.append(child)
    return files

lists = []
for filepath in getfilelist(Folderpath):
    if filepath.endswith("docx") and not filepath.startswith('~$'):
        lists.append(filepath)

for filepath in lists:
    try:
        # 数据源              
        tfh_str = ''
        jz_count = 0
        dkh_array = []
        dlh_array = []
        dkm_array = []
        f1 = open(os.path.dirname(filepath) + "/" + "宗地.txt","w+",encoding="utf-8")
        f2 = open(os.path.dirname(filepath) + "/" + "界址点坐标文件.txt","w+",encoding="utf-8")
        f1.writelines('[属性描述]\n坐标系=2000国家大地坐标系\n几度分带=3\n投影类型=高斯克吕格\n计量单位=米\n带号=37\n精度=0.001\n转换参数=,,,,,,\n[地块坐标]\n')
        f2.writelines('[属性描述]\n坐标系=2000国家大地坐标系\n几度分带=3\n投影类型=高斯克吕格\n计量单位=米\n带号=37\n精度=0.001\n转换参数=,,,,,,\n[地块坐标]\n')
        doc = Document(filepath) 
        for p in doc.paragraphs:
            if p.text == '界址点成果表':
                    jz_count = jz_count+1
            if '地块号' in p.text:
                dkh_array.append(p.text.split()[1])
                dlh_array.append(p.text.split()[3])
                dkm_array.append(p.text.split('：')[-1])  
        tfh_str = doc.tables[0].cell(5,2).text
        for i in range(4,len(doc.tables)-jz_count):
            table = doc.tables[i]
            dk_mc = '地块'+str(i-3)  
            dk_count = int((len(table.rows)-2)/2)
            area = table.cell(len(table.rows)-1,0).text.split()[2]
            area_gq = format(float(area)/10000,'.4f')
            f1.writelines(str(dk_count)+','+str(area_gq)+','+str(dkh_array[i-4])+','+str(dlh_array[i-4])+','+str(dkm_array[i-4])+','+','+dk_mc+'@'+'\n')
            for j in range(0,dk_count):
                coo = (j+1)*2
                f1.writelines(table.cell(coo,0).text+','+'1'+','+table.cell(coo,1).text+','+table.cell(coo,2).text+'\n')
        for m in range(len(doc.tables)-jz_count,len(doc.tables)):
            table = doc.tables[m]
            dk = str(m-(len(doc.tables)-jz_count-1))
            dk_mc = '地块'+dk
            dk_count = int((len(table.rows)-2)/2)
            area = table.cell(len(table.rows)-1,0).text.split()[2]
            area_gq = format(float(area)/10000,'.4f')
            f2.writelines(str(dk_count)+','+str(area_gq)+','+dk+','+dk_mc+','+'面'+','+tfh_str+',,,@'+'\n')
            for n in range(0,dk_count):
                coo = (n+1)*2
                f2.writelines(table.cell(coo,0).text+','+'1'+','+table.cell(coo,1).text+','+table.cell(coo,2).text+'\n')
        f1.close
        f2.close
    except:
        print('问题路径---',filepath )
        continue
messagebox.showinfo('提示','地块界址点坐标txt完成！')