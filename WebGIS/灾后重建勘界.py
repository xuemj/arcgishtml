import os
from tkinter import filedialog
from docx import Document
import tkinter as tk
from tkinter import messagebox
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import PatternFill,Border, Side, Alignment, Font
alignment=Alignment(horizontal='center',vertical='center',wrap_text=True)


font1 = Font(
    name="微软雅黑",   # 字体
    size=18,         # 字体大小
    bold=True,       # 是否加粗，True/False
)

font2 = Font(
    name="微软雅黑",   # 字体
    size=11,         # 字体大小
)

font3 = Font(
    name="微软雅黑",   # 字体
    size=12,         # 字体大小
    bold=True,       # 是否加粗，True/False
)

font4 = Font(
    name="宋体",   # 字体
    size=11,         # 字体大小
)

border = Border(left=Side(border_style='thin', color='000000'),
                right=Side(border_style='thin', color='000000'),
                top=Side(border_style='thin', color='000000'),
                bottom=Side(border_style='thin', color='000000'))

root = tk.Tk()
root.withdraw()

messagebox.showinfo('提示','请选择文件夹！')

Folderpath = filedialog.askdirectory() #获得选择好的文件夹

# 获取filepath文件夹下的所有的文件
def getfilelist(filepath):
    filelist =  os.listdir(filepath)
    files = []
    for i in range(len(filelist)):
        child = os.path.join('%s/%s'%(filepath, filelist[i]))
        if os.path.isdir(child):
            files.extend(getfilelist(child))
        else:
            files.append(child)
    return files

lists = []
for filepath in getfilelist(Folderpath):
    if filepath.endswith("docx") and not filepath.startswith('~$'):
        lists.append(filepath)

for filepath in lists:
    try:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.row_dimensions[1].height = 38.85
        worksheet.row_dimensions[2].height = 30
        worksheet.row_dimensions[3].height = 31.35
        worksheet.row_dimensions[4].height = 40
        worksheet.row_dimensions[5].height = 23.45
        worksheet.column_dimensions['A'].width = 6
        worksheet.column_dimensions['B'].width = 23
        worksheet.column_dimensions['C'].width = 10
        worksheet.column_dimensions['D'].width = 7
        worksheet.column_dimensions['E'].width = 13
        worksheet.column_dimensions['F'].width = 13
        worksheet.column_dimensions['G'].width = 13
        worksheet.column_dimensions['H'].width = 13
        worksheet.column_dimensions['I'].width = 13
        worksheet.column_dimensions['J'].width = 10
        worksheet.column_dimensions['K'].width = 13
        worksheet.column_dimensions['L'].width = 13
        worksheet.column_dimensions['M'].width = 13
        worksheet.column_dimensions['N'].width = 13
        worksheet.column_dimensions['O'].width = 13
        worksheet.column_dimensions['P'].width = 13
        worksheet.column_dimensions['Q'].width = 13
        worksheet.column_dimensions['R'].width = 13
        worksheet.column_dimensions['S'].width = 13
        worksheet.column_dimensions['T'].width = 13
        worksheet.column_dimensions['U'].width = 13
        worksheet.column_dimensions['V'].width = 13
        worksheet.column_dimensions['W'].width = 13
        worksheet.column_dimensions['X'].width = 13
        worksheet.column_dimensions['Y'].width = 13
        worksheet.column_dimensions['Z'].width = 13
        worksheet.cell(1,1).value = '土地分类面积表'
        worksheet.cell(1,1).alignment = alignment
        worksheet.cell(1,1).font = font1
        worksheet.cell(1,1).border = border
        worksheet.cell(2,1).value = '单位：公顷'
        worksheet.cell(2,1).alignment = Alignment(vertical='center',wrap_text=True)
        worksheet.cell(2,1).font = font2
        worksheet.cell(2,1).border = border
        worksheet.cell(3,1).value = '序号'
        worksheet.cell(3,1).alignment = alignment
        worksheet.cell(3,1).font = font3
        worksheet.cell(3,1).border = border
        worksheet.cell(3,2).value = '所在地'
        worksheet.cell(3,2).alignment = alignment
        worksheet.cell(3,2).font = font3
        worksheet.cell(3,2).border = border
        worksheet.cell(3,3).value = '权属单位'
        worksheet.cell(3,3).alignment = alignment
        worksheet.cell(3,3).font = font3
        worksheet.cell(3,3).border = border
        worksheet.cell(3,4).value = '村组\n(社区)'
        worksheet.cell(3,4).alignment = alignment
        worksheet.cell(3,4).font = font3
        worksheet.cell(3,4).border = border
        worksheet.cell(3,5).value = '权属类别'
        worksheet.cell(3,5).alignment = alignment
        worksheet.cell(3,5).font = font3
        worksheet.cell(3,5).border = border
        worksheet.cell(3,6).value = '征地前人均\n耕地(亩/人)'
        worksheet.cell(3,6).font = font3
        worksheet.cell(3,6).alignment = alignment
        worksheet.cell(3,6).border = border
        worksheet.cell(3,7).value = '征地后人均\n耕地(亩/人)'
        worksheet.cell(3,7).font = font3
        worksheet.cell(3,7).alignment = alignment
        worksheet.cell(3,7).border = border
        worksheet.cell(3,8).value = '拟安置农业\n人口数(人)'
        worksheet.cell(3,8).font = font3
        worksheet.cell(3,8).alignment = alignment
        worksheet.cell(3,8).border = border
        worksheet.cell(3,9).value = '拟安置劳动\n力人数(人)'
        worksheet.cell(3,9).font = font3
        worksheet.cell(3,9).alignment = alignment
        worksheet.cell(3,9).border = border
        worksheet.cell(3,10).value = '面积总计'
        worksheet.cell(3,10).alignment = alignment
        worksheet.cell(3,10).font = font3
        worksheet.cell(3,10).border = border
        worksheet.merge_cells(start_row=3,end_row=4,start_column=1,end_column=1)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=2,end_column=2)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=3,end_column=3)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=4,end_column=4)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=5,end_column=5)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=6,end_column=6)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=7,end_column=7)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=8,end_column=8)

        worksheet.merge_cells(start_row=3,end_row=4,start_column=9,end_column=9)
        worksheet.merge_cells(start_row=3,end_row=4,start_column=10,end_column=10)

        doc = Document(filepath) 
        table = doc.tables[2]
        table1 = doc.tables[3]
        count = len(table1.columns)
        gd_name_arr = []
        yd_name_arr = []
        ld_name_arr = []
        jt_name_arr = []
        qt_name_arr = []
        gd_value_arr = []
        yd_value_arr = []
        ld_value_arr = []
        jt_value_arr = []
        qt_value_arr = []
        bj_str = ''
        for i in range(1,count-6):
            if table1.cell(2,i).text == '耕地':
                bj_str = 'gd'
                continue
            if table1.cell(2,i).text == '种植园用地':
                bj_str = 'yd'
                continue
            if table1.cell(2,i).text == '林地': 
                bj_str = 'ld'
                continue
            if table1.cell(2,i).text == '交通用地':
                bj_str = 'jt'
                continue
            if table1.cell(2,i).text == '其他土地':
                bj_str = 'qt'
                continue
            if bj_str =='gd':
                if table1.cell(3,i).text != '':
                    gd_name_arr.append(table1.cell(2,i).text)
                    gd_value_arr.append(table1.cell(3,i).text)
            if bj_str =='yd':
                if table1.cell(3,i).text != '':
                    yd_name_arr.append(table1.cell(2,i).text)
                    yd_value_arr.append(table1.cell(3,i).text)
            if bj_str == 'ld':
                if table1.cell(3,i).text != '':
                    ld_name_arr.append(table1.cell(2,i).text)
                    ld_value_arr.append(table1.cell(3,i).text)
            if bj_str == 'jt':
                if table1.cell(3,i).text != '':
                    jt_name_arr.append(table1.cell(2,i).text)
                    jt_value_arr.append(table1.cell(3,i).text)
            if bj_str == 'qt':
                if table1.cell(3,i).text != '':
                    qt_name_arr.append(table1.cell(2,i).text)
                    qt_value_arr.append(table1.cell(3,i).text)
        worksheet.cell(5,1).value = '1'
        worksheet.cell(5,1).alignment = alignment
        worksheet.cell(5,1).font = font4
        worksheet.cell(5,1).border = border
        worksheet.cell(5,2).value = '济源市' + '-' +'济源市' + '-' + table.cell(3,0).text
        worksheet.cell(5,2).alignment = alignment
        worksheet.cell(5,2).font = font4
        worksheet.cell(5,2).border = border
        worksheet.cell(5,3).value = table1.cell(3,0).text
        worksheet.cell(5,3).alignment = alignment
        worksheet.cell(5,3).font = font4
        worksheet.cell(5,3).border = border
        worksheet.cell(5,4).border = border
        worksheet.cell(5,5).value = '使用集体土地'
        worksheet.cell(5,5).alignment = alignment
        worksheet.cell(5,5).font = font4
        worksheet.cell(5,5).border = border
        worksheet.cell(5,6).value = '1'
        worksheet.cell(5,6).alignment = alignment
        worksheet.cell(5,6).font = font4
        worksheet.cell(5,6).border = border
        worksheet.cell(5,7).value = '1'
        worksheet.cell(5,7).alignment = alignment
        worksheet.cell(5,7).font = font4
        worksheet.cell(5,7).border = border
        worksheet.cell(5,8).value = '1'
        worksheet.cell(5,8).alignment = alignment
        worksheet.cell(5,8).font = font4
        worksheet.cell(5,8).border = border
        worksheet.cell(5,9).value = '1'
        worksheet.cell(5,9).alignment = alignment
        worksheet.cell(5,9).font = font4
        worksheet.cell(5,9).border = border
        worksheet.cell(5,10).value = table1.cell(3,count-2).text
        worksheet.cell(5,10).alignment = alignment
        worksheet.cell(5,10).font = font4
        worksheet.cell(5,10).border = border
        cl = 10
        if len(gd_value_arr) > 0 :
            gd_xj = 0
            if len(qt_name_arr) > 0:
                for i in range(len(qt_name_arr)) :
                    if qt_name_arr[i] == '田坎' :
                        gd_xj = float(gd_value_arr[0])+float(qt_value_arr[i])
                    else :
                        gd_xj = float(gd_value_arr[0])
            else :
                gd_xj = float(gd_value_arr[0])
            gd_xj = format(gd_xj,'.4f')
            cllll = cl
            worksheet.cell(3,cl+1).value = '耕地'
            worksheet.cell(4,cl+1).value ='耕地小计'
            worksheet.cell(4,cl+2).value = gd_name_arr[0]
            worksheet.cell(5,cl+1).value = str(gd_xj)
            worksheet.cell(5,cl+2).value = str(gd_xj)
            worksheet.cell(3,cl+1).alignment = alignment
            worksheet.cell(3,cl+1).font = font3
            worksheet.cell(3,cl+1).border = border
            worksheet.cell(4,cl+1).alignment = alignment
            worksheet.cell(4,cl+1).font = font3
            worksheet.cell(4,cl+1).border = border
            worksheet.cell(4,cl+2).alignment = alignment
            worksheet.cell(4,cl+2).font = font3
            worksheet.cell(4,cl+2).border = border
            worksheet.cell(5,cl+1).alignment = alignment
            worksheet.cell(5,cl+1).font = font4
            worksheet.cell(5,cl+1).border = border
            worksheet.cell(5,cl+2).alignment = alignment
            worksheet.cell(5,cl+2).font = font4
            worksheet.cell(5,cl+2).border = border
            cl = cl + 2
            worksheet.merge_cells(start_row=3,end_row=3,start_column=cllll+1,end_column=cl)
        if len(yd_value_arr) > 0 :
            cllll = cl
            worksheet.cell(3,cl+1).value = '种植园用地'
            worksheet.cell(4,cl+1).value ='种植园用地小计'
            worksheet.cell(4,cl+2).value = yd_name_arr[0]
            worksheet.cell(5,cl+1).value = yd_value_arr[0]
            worksheet.cell(5,cl+2).value = yd_value_arr[0]
            worksheet.cell(3,cl+1).alignment = alignment
            worksheet.cell(3,cl+1).font = font3
            worksheet.cell(3,cl+1).border = border
            worksheet.cell(4,cl+1).alignment = alignment
            worksheet.cell(4,cl+1).font = font3
            worksheet.cell(4,cl+1).border = border
            worksheet.cell(4,cl+2).alignment = alignment
            worksheet.cell(4,cl+2).font = font3
            worksheet.cell(4,cl+2).border = border
            worksheet.cell(5,cl+1).alignment = alignment
            worksheet.cell(5,cl+1).font = font4
            worksheet.cell(5,cl+1).border = border
            worksheet.cell(5,cl+2).alignment = alignment
            worksheet.cell(5,cl+2).font = font4
            worksheet.cell(5,cl+2).border = border
            cl = cl + 2
            worksheet.merge_cells(start_row=3,end_row=3,start_column=cllll+1,end_column=cl)
        if len(ld_value_arr) > 0 :
            cllll = cl
            worksheet.cell(3,cl+1).value = '林地'
            worksheet.cell(4,cl+1).value = '林地小计'
            worksheet.cell(3,cl+1).alignment = alignment
            worksheet.cell(3,cl+1).font = font3
            worksheet.cell(3,cl+1).border = border
            worksheet.cell(4,cl+1).alignment = alignment
            worksheet.cell(4,cl+1).font = font3
            worksheet.cell(4,cl+1).border = border
            ld_xj = 0
            for a in range(len(ld_name_arr)) :
                worksheet.cell(4,cl+2+a).value = ld_name_arr[a]
                worksheet.cell(5,cl+2+a).value = ld_value_arr[a]
                worksheet.cell(4,cl+2+a).alignment = alignment
                worksheet.cell(4,cl+2+a).font = font3
                worksheet.cell(4,cl+2+a).border = border
                worksheet.cell(5,cl+2+a).alignment = alignment
                worksheet.cell(5,cl+2+a).font = font4
                worksheet.cell(5,cl+2+a).border = border
                ld_xj += float(ld_value_arr[a])
            ld_xj = format(ld_xj,'.4f')
            worksheet.cell(5,cl+1).value = str(ld_xj)
            worksheet.cell(5,cl+1).alignment = alignment
            worksheet.cell(5,cl+1).font = font4
            worksheet.cell(5,cl+1).border = border
            cl = cl + 1 + len(ld_name_arr)
            worksheet.merge_cells(start_row=3,end_row=3,start_column=cllll+1,end_column=cl)
        if len(jt_value_arr) > 0 :
            cllll = cl
            worksheet.cell(3,cl+1).value = '交通用地'
            worksheet.cell(4,cl+1).value ='交通用地小计'
            worksheet.cell(4,cl+2).value = jt_name_arr[0]
            worksheet.cell(5,cl+1).value = jt_value_arr[0]
            worksheet.cell(5,cl+2).value = jt_value_arr[0]
            worksheet.cell(3,cl+1).alignment = alignment
            worksheet.cell(3,cl+1).font = font3
            worksheet.cell(3,cl+1).border = border
            worksheet.cell(4,cl+1).alignment = alignment
            worksheet.cell(4,cl+1).font = font3
            worksheet.cell(4,cl+1).border = border
            worksheet.cell(4,cl+2).alignment = alignment
            worksheet.cell(4,cl+2).font = font3
            worksheet.cell(4,cl+2).border = border
            worksheet.cell(5,cl+1).alignment = alignment
            worksheet.cell(5,cl+1).font = font4
            worksheet.cell(5,cl+1).border = border
            worksheet.cell(5,cl+2).alignment = alignment
            worksheet.cell(5,cl+2).font = font4
            worksheet.cell(5,cl+2).border = border
            cl = cl + 2
            worksheet.merge_cells(start_row=3,end_row=3,start_column=cllll+1,end_column=cl)
        if len(qt_value_arr) > 0 :
            if qt_name_arr[0] != '田坎' :
                cllll = cl
                worksheet.cell(3,cl+1).value = '其他土地'
                worksheet.cell(4,cl+1).value = '其他小计'
                worksheet.cell(4,cl+2).value = qt_name_arr[0]
                worksheet.cell(5,cl+1).value = qt_value_arr[0]
                worksheet.cell(5,cl+2).value = qt_value_arr[0]
                worksheet.cell(3,cl+1).alignment = alignment
                worksheet.cell(3,cl+1).font = font3
                worksheet.cell(3,cl+1).border = border
                worksheet.cell(4,cl+1).alignment = alignment
                worksheet.cell(4,cl+1).font = font3
                worksheet.cell(4,cl+1).border = border
                worksheet.cell(4,cl+2).alignment = alignment
                worksheet.cell(4,cl+2).font = font3
                worksheet.cell(4,cl+2).border = border
                worksheet.cell(5,cl+1).alignment = alignment
                worksheet.cell(5,cl+1).font = font4
                worksheet.cell(5,cl+1).border = border
                worksheet.cell(5,cl+2).alignment = alignment
                worksheet.cell(5,cl+2).font = font4
                worksheet.cell(5,cl+2).border = border
                cl = cl + 2
                worksheet.merge_cells(start_row=3,end_row=3,start_column=cllll+1,end_column=cl)
        worksheet.merge_cells(start_row=1,end_row=1,start_column=1,end_column=cl)
        worksheet.merge_cells(start_row=2,end_row=2,start_column=1,end_column=cl)
        workbook.save(os.path.dirname(filepath) + '/' + '宗地.xlsx')
    except:
        print('问题路径---',filepath )
        continue
messagebox.showinfo('提示','勘界面积EXCEL完成！')