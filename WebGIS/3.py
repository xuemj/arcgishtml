import openpyxl
import tkinter as tk
from tkinter import filedialog
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = tk.Tk()
root.withdraw()
selected_file_path = filedialog.askopenfilename()  # 使用askopenfilename函数选择单个文件

df = pd.read_excel(selected_file_path)
wk = openpyxl.load_workbook(selected_file_path)
wk_name = wk.sheetnames
wk_sheet = wk[wk_name[0]]
doc = Document('样地调查表.docx')
for x in range(1, 33):
    table1 = doc.add_table(rows=6, cols=6, style='Table Grid')
    table1.cell(0, 0).text = '样 地 信 息'
    table1.cell(0, 0).merge(table1.cell(0, 5))
    table1.cell(1, 0).text = '样地编号'
    table1.cell(1, 0).merge(table1.cell(2, 0))
    table1.cell(1, 1).text = str(x + 1)
    table1.cell(1, 1).merge(table1.cell(2, 1))
    table1.cell(3, 0).text = '样地面积'
    table1.cell(3, 0).merge(table1.cell(4, 0))
    table1.cell(3, 1).merge(table1.cell(4, 1))
    table1.cell(5, 0).merge(table1.cell(5, 5))
    table1.cell(1, 2).text = '样地位置'
    table1.cell(1, 2).merge(table1.cell(4, 2))
    paragraph = doc.add_paragraph("调 查 信 息")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table2 = doc.add_table(rows=48, cols=12, style='Table Grid')
    table2.cell(0,0).merge(table2.cell(0,11))
    table2.cell(1, 0).text = '坑\
                                        编\
                                        号'
    table2.cell(1, 0).merge(table2.cell(2, 0))
    table2.cell(1, 1).text = '合格数量'
    table2.cell(1, 1).merge(table2.cell(1, 2))
    table2.cell(1, 3).text = '成活苗木'
    table2.cell(1, 3).merge(table2.cell(1, 5))
    table2.cell(2, 1).text = '鱼鳞坑'
    table2.cell(2, 2).text = '育林板'
    table2.cell(2, 3).text = '侧柏'
    table2.cell(2, 4).text = '栓皮栎'
    table2.cell(2, 5).text = '元宝枫'
    table2.cell(1, 6).text = '坑\
                                        编\
                                        号'
    table2.cell(1, 6).merge(table2.cell(2, 6))
    table2.cell(1, 7).text = '合格数量'
    table2.cell(1, 7).merge(table2.cell(1, 8))
    table2.cell(1, 9).text = '成活苗木'
    table2.cell(1, 9).merge(table2.cell(1, 11))
    table2.cell(2, 7).text = '鱼鳞坑'
    table2.cell(2, 8).text = '育林板'
    table2.cell(2, 9).text = '侧柏'
    table2.cell(2, 10).text = '栓皮栎'
    table2.cell(2, 11).text = '元宝枫'
    table2.cell(43, 0).merge(table2.cell(43, 11))
    table2.cell(43, 2).text = '核 查 结 论'
    table2.cell(44, 0).merge(table2.cell(45, 1))
    table2.cell(44, 0).text = '样地合计'
    table2.cell(45, 0).merge(table2.cell(47, 1))
    table2.cell(45, 1).text = '整地'
    table2.cell(44, 2).merge(table2.cell(44, 3))
    table2.cell(44, 2).text = '类型'
    table2.cell(45, 2).merge(table2.cell(46, 3))
    table2.cell(45, 2).text = '鱼鳞坑'
    table2.cell(47, 2).merge(table2.cell(47, 3))
    table2.cell(47, 2).text = '育林板'
    table2.cell(44, 4).merge(table2.cell(44, 5))
    table2.cell(45, 4).merge(table2.cell(46, 5))
    table2.cell(47, 4).merge(table2.cell(47, 5))
    table2.cell(44, 4).text = '合格数量'
    table2.cell(45, 2).merge(table2.cell(46, 3))
    table2.cell(47, 2).merge(table2.cell(47, 3))
    table2.cell(44, 6).merge(table2.cell(44, 7))
    table2.cell(44, 6).text = '样地合计'
    table2.cell(45, 6).merge(table2.cell(47, 7))
    table2.cell(45, 6).text = '苗木'
    table2.cell(44, 8).merge(table2.cell(44, 9))
    table2.cell(44, 8).text = '树种'
    table2.cell(45, 8).merge(table2.cell(45, 9))
    table2.cell(45, 8).text = '侧柏'
    table2.cell(46, 8).merge(table2.cell(46, 9))
    table2.cell(46, 8).text = '栓皮栎'
    table2.cell(47, 8).merge(table2.cell(47, 9))
    table2.cell(47, 8).text = '元宝枫'
    table2.cell(44, 10).merge(table2.cell(44, 11))
    table2.cell(44, 10).text = '成活数量'
    table2.cell(45, 10).merge(table2.cell(45, 11))
    table2.cell(46, 10).merge(table2.cell(46, 11))
    table2.cell(47, 10).merge(table2.cell(47, 11))
    paragraph = doc.add_paragraph("")
print(len(doc.tables))
count = (int)(len(doc.tables) / 2)
for y in range(count):
    b = 2 * y
    a = 2 * y + 1
    print('-------------------', b)
    table = doc.tables[b]
    n = int(table.cell(1, 1).text)
    table3 = doc.tables[a]
    num = 1
    z = 1
    q = 0
    w = 0
    e = 0
    r = 0
    for l in df.values[0:]:
        if l[8] == n:
            data = (l[14], l[15], l[16], l[17])
            print(l[8], num, data)
            if l[14] == 1:
                q += 1
            print(q)
            if l[15] == 1:
                w += 1
            print(w)
            if l[16] == 1:
                r += 1
            elif l[16] == 2:
                r += 2
            print(r)
            if l[17] == 1:
                e += 1
            elif l[17] == 2:
                e += 2
            print(r)
            if num <= 40:
                table3.cell(num + 2, 0).text = str(num)
                table3.cell(num + 2, 1).text = str(data[0])
                table3.cell(num + 2, 2).text = str(data[1])
                table3.cell(num + 2, 3).text = str(data[2])
                table3.cell(num + 2, 5).text = str(data[3])
            if num > 40 and num <= 80:
                table3.cell(z + 2, 6).text = str(num)
                table3.cell(z + 2, 7).text = str(data[0])
                table3.cell(z + 2, 8).text = str(data[1])
                table3.cell(z + 2, 9).text = str(data[2])
                table3.cell(z + 2, 11).text = str(data[3])
                z += 1
            table3.cell(45, 4).text = str(q)
            table3.cell(47, 4).text = str(w)
            table3.cell(45, 10).text = str(r)
            table3.cell(47, 10).text = str(e)
            num += 1
doc.save('样地调查表.docx')
